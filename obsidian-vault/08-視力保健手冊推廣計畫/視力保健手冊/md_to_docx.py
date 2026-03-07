# -*- coding: utf-8 -*-
"""
V19 Markdown → Word (.docx) 轉換腳本
使用 python-docx 進行精細排版，逐行解析 Markdown
"""

import re
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = r"c:\Users\w7\Dropbox\PC (2)\Desktop\uncleglasses\optometry-notes\projects\視力保健手冊\完整手冊_V19.md"
OUTPUT = r"c:\Users\w7\Dropbox\PC (2)\Desktop\uncleglasses\optometry-notes\projects\視力保健手冊\完整手冊_V19.docx"

doc = Document()

# ─── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ─── Style helpers ──────────────────────────────────────────────────────────────
styles = doc.styles

def set_font(run, size=12, bold=False, color=None, name='Noto Serif TC'):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # CJK font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.insert(0, rFonts)

def para_space(p, before=0, after=4):
    pPr = p._p.get_or_add_pPr()
    pPr_tag = OxmlElement('w:spacing')
    pPr_tag.set(qn('w:before'), str(before*20))
    pPr_tag.set(qn('w:after'), str(after*20))
    pPr.append(pPr_tag)

def shade_cell(cell, fill='EFF3FB'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_heading(doc, text, level, emoji_ok=True):
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
    style = style_map.get(level, 'Heading 4')
    p = doc.add_heading(text, level=level)
    p.style = doc.styles[style]
    # color
    colors = {1:(0x1A,0x56,0xAD), 2:(0x17,0x59,0x6E), 3:(0x2E,0x74,0x54), 4:(0x57,0x35,0x6D)}
    for run in p.runs:
        run.font.color.rgb = RGBColor(*colors.get(level,(0,0,0)))
        set_font(run, size=[18,16,14,13][level-1], bold=True, name='Noto Serif TC')
    para_space(p, before=[12,8,6,4][level-1], after=[6,4,3,2][level-1])
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    # handle inline bold (**...**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_font(run, size=11, bold=True)
        else:
            run = p.add_run(part)
            set_font(run, size=11)
    para_space(p, before=0, after=3)
    return p

def add_callout(doc, kind, text):
    """Render [!TIP], [!NOTE], [!IMPORTANT], [!WARNING] as shaded paragraph."""
    config = {
        'TIP':       ('💡 速讀提示', 'D1EAD1', (0x1D, 0x6A, 0x1D)),
        'NOTE':      ('📝 注意',      'FFF3CD', (0x99, 0x66, 0x00)),
        'IMPORTANT': ('🔑 重要',      'D6E4FF', (0x0D, 0x3C, 0x8C)),
        'WARNING':   ('⚠️ 警告',     'FFE4E4', (0x99, 0x00, 0x00)),
        'CAUTION':   ('🚨 注意事項',  'F9D6FF', (0x5C, 0x00, 0x7E)),
    }
    label, fill, txt_color = config.get(kind, ('📌', 'EFEFEF', (0,0,0)))
    p = doc.add_paragraph()
    # Label
    run = p.add_run(f"{label}  ")
    set_font(run, size=10, bold=True, color=txt_color)
    # Content
    run2 = p.add_run(text)
    set_font(run2, size=10, color=txt_color)
    # Shade
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)
    # Indent
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:right'), '360')
    pPr.append(ind)
    para_space(p, before=3, after=3)
    return p

def render_inline(text):
    """Strip markdown inline formatting for table cells."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text

def add_table_from_rows(doc, rows):
    """rows is list of list[str]. First row = header."""
    if not rows or not rows[0]:
        return
    col_count = max(len(r) for r in rows)
    # Normalize
    rows = [r + ['']*(col_count-len(r)) for r in rows]
    t = doc.add_table(rows=len(rows), cols=col_count)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = t.cell(i, j)
            cell_text = render_inline(cell_text.strip())
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                set_font(run, size=9.5, bold=(i==0))
            if i == 0:
                shade_cell(cell, fill='1A56AD')
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            elif i % 2 == 0:
                shade_cell(cell, fill='EFF3FB')
    doc.add_paragraph()

def add_code_block(doc, lines):
    """Render code (including mermaid) as shaded block."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    pPr.append(ind)
    text = '\n'.join(lines)
    if text.strip().startswith(('graph', 'gantt', 'sequenceDiagram', 'pie', 'flowchart')):
        run = p.add_run('📊 [Mermaid 圖表 — 請以 Markdown 預覽工具查看]\n' + text)
    else:
        run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Consolas')
    rPr.insert(0, rFonts)
    para_space(p, before=3, after=3)

# ─── Title Page ─────────────────────────────────────────────────────────────────
def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n高級中等以下學校學生視力保健手冊')
    set_font(run, size=24, bold=True, color=(0x1A,0x56,0xAD))
    para_space(p, before=0, after=6)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('以實證為基礎的視力保健指引')
    set_font(run2, size=16, bold=False, color=(0x17,0x59,0x6E))
    para_space(p2, before=0, after=4)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('臺灣視光視力保健學會\n2026 年 3 月（第二版，V19）')
    set_font(run3, size=12, color=(0x44,0x44,0x44))
    para_space(p3, before=6, after=0)

    doc.add_page_break()

add_title_page(doc)

# ─── Parse Markdown ─────────────────────────────────────────────────────────────
with open(INPUT, 'r', encoding='utf-8') as f:
    raw = f.read()

lines = raw.splitlines()

# Skip YAML frontmatter
start = 0
if lines and lines[0].strip() == '---':
    for i in range(1, len(lines)):
        if lines[i].strip() == '---':
            start = i + 1
            break
lines = lines[start:]

# State machine
i = 0
table_buf = []
code_buf = []
in_code = False
code_lang = ''
callout_buf = []
callout_kind = ''

def flush_table():
    global table_buf
    if not table_buf:
        return
    rows = []
    for row_line in table_buf:
        # separator line
        if re.match(r'^\|?\s*[-: |]+\s*\|?$', row_line):
            continue
        cells = [c.strip() for c in row_line.strip('|').split('|')]
        rows.append(cells)
    add_table_from_rows(doc, rows)
    table_buf = []

def flush_callout():
    global callout_buf, callout_kind
    if callout_buf:
        text = ' '.join(callout_buf).strip()
        add_callout(doc, callout_kind, text)
    callout_buf = []
    callout_kind = ''

total = len(lines)
while i < total:
    line = lines[i]
    stripped = line.strip()

    # ── Code block open/close
    if stripped.startswith('```'):
        if in_code:
            # closing
            add_code_block(doc, code_buf)
            code_buf = []
            in_code = False
        else:
            flush_table()
            flush_callout()
            in_code = True
            code_lang = stripped[3:].strip()
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # ── Skip \newpage
    if stripped == r'\newpage' or stripped == '\\newpage':
        flush_table()
        flush_callout()
        doc.add_page_break()
        i += 1
        continue

    # ── Callout block  > [!KIND]
    callout_match = re.match(r'^>\s*\[!(TIP|NOTE|IMPORTANT|WARNING|CAUTION)\]', stripped)
    if callout_match:
        flush_table()
        flush_callout()
        callout_kind = callout_match.group(1)
        callout_buf = []
        i += 1
        continue

    if callout_kind and stripped.startswith('>'):
        content = stripped.lstrip('>').strip()
        # remove leading ** bold markers but keep text
        content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
        callout_buf.append(content)
        i += 1
        continue

    if callout_kind and not stripped.startswith('>'):
        flush_callout()
        # fall through to process current line

    # ── Horizontal rule
    if re.match(r'^[-*_]{3,}\s*$', stripped):
        flush_table()
        p = doc.add_paragraph('─' * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            run.font.size = Pt(8)
        para_space(p, before=2, after=2)
        i += 1
        continue

    # ── Headings
    h_match = re.match(r'^(#{1,6})\s+(.+)', stripped)
    if h_match:
        flush_table()
        level = len(h_match.group(1))
        text = h_match.group(2).strip()
        add_heading(doc, text, min(level, 4))
        i += 1
        continue

    # ── Table rows
    if stripped.startswith('|') or (table_buf and stripped.startswith('|')):
        table_buf.append(stripped)
        i += 1
        continue
    else:
        flush_table()

    # ── Blank line
    if not stripped:
        i += 1
        continue

    # ── Bullet list
    bullet_match = re.match(r'^[-*+]\s+(.*)', stripped)
    if bullet_match:
        p = doc.add_paragraph(style='List Bullet')
        content = bullet_match.group(1)
        parts = re.split(r'(\*\*.*?\*\*|`.*?`)', content)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_font(run, size=10.5, bold=True)
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
            else:
                run = p.add_run(part)
                set_font(run, size=10.5)
        para_space(p, before=0, after=2)
        i += 1
        continue

    # ── Ordered list
    ol_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
    if ol_match:
        p = doc.add_paragraph(style='List Number')
        content = ol_match.group(2)
        parts = re.split(r'(\*\*.*?\*\*)', content)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                set_font(run, size=10.5, bold=True)
            else:
                run = p.add_run(part)
                set_font(run, size=10.5)
        para_space(p, before=0, after=2)
        i += 1
        continue

    # ── Blockquote (non-callout)
    if stripped.startswith('>'):
        flush_table()
        content = stripped.lstrip('>').strip()
        content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '480')
        pPr.append(ind)
        run = p.add_run(content)
        set_font(run, size=10.5, color=(0x33,0x33,0x77))
        run.italic = True
        para_space(p, before=2, after=2)
        i += 1
        continue

    # ── Normal paragraph
    add_body(doc, stripped)
    i += 1

flush_table()
flush_callout()

# ─── Save ────────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"✅ Done! Saved to: {OUTPUT}")
